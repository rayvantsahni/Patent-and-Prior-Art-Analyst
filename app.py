import streamlit as st
import sys
import os
import io  # Needed for in-memory file creation
from datetime import datetime  # Needed for smart filenames
from docx import Document  # Needed for .docx creation

# --- Path Setup ---
# Add the 'src' directory to the Python path
# This allows us to import from 'src.backend'
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

try:
    from src.backend import agent
    from src.backend.rate_limiter import SimpleRateLimiter
except ImportError:
    st.error("Error: Could not import the backend agent. Make sure 'src/backend/agent.py' exists.")
    st.stop()


# --- Internal Helper Functions ---
def _generate_filename_prefix():
    """Generates a clean, timestamped filename prefix."""
    now = datetime.now()
    return f"patent_analysis_{now.strftime('%Y%m%d_%H%M%S')}"


def _create_docx(content):
    """Creates an in-memory .docx file from the report text."""
    doc = Document()
    doc.add_heading('Patent Analysis Report', 0)

    # Add the full report as paragraphs
    # This preserves line breaks
    for line in content.split('\n'):
        doc.add_paragraph(line)

    # Save to an in-memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --- Helper for displaying artifacts ---
def _display_artifacts_side_by_side(base_artifacts, novel_artifacts):
    """Helper to neatly display keywords, CPCs, and HyDE abstract side by side."""

    # Technical Keywords & Synonyms
    st.markdown("<h3 style='text-align: center;'>Technical Keywords & Synonyms</h3>", unsafe_allow_html=True)
    st.markdown("> Use these terms for your own research on Google Patents.")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Base Technology Search**")
        base_keywords = base_artifacts.get('technical_keywords', [])
        tags_html = "".join([
            f"<span style='background-color: #eee; border-radius: 5px; padding: 3px 8px; margin: 3px; display: inline-block;'>{kw}</span>"
            for kw in base_keywords])
        st.markdown(tags_html, unsafe_allow_html=True)

    with col2:
        st.markdown("**Novel Features Search**")
        novel_keywords = novel_artifacts.get('technical_keywords', [])
        tags_html = "".join([
            f"<span style='background-color: #eee; border-radius: 5px; padding: 3px 8px; margin: 3px; display: inline-block;'>{kw}</span>"
            for kw in novel_keywords])
        st.markdown(tags_html, unsafe_allow_html=True)

    st.divider()

    # Relevant CPC Codes
    st.markdown("<h3 style='text-align: center;'>Relevant CPC Codes</h3>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 0.9em; color: #666;'>CPC stands for 'Cooperative Patent Classification'. This is a professional, expert-level 'tag' used to categorize patents. You can paste these codes directly into Google Patents to find all patents in that specific category.</p>", unsafe_allow_html=True)
    st.markdown("> These are the 'expert-level' classifications for this technology.")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Base Technology Search**")
        base_cpcs = base_artifacts.get('cpc_codes', [])
        tags_html = "".join([
            f"<span style='background-color: #e0f2fe; border-radius: 5px; padding: 3px 8px; margin: 3px; display: inline-block;'>{cpc}</span>"
            for cpc in base_cpcs])
        st.markdown(tags_html, unsafe_allow_html=True)

    with col2:
        st.markdown("**Novel Features Search**")
        novel_cpcs = novel_artifacts.get('cpc_codes', [])
        tags_html = "".join([
            f"<span style='background-color: #e0f2fe; border-radius: 5px; padding: 3px 8px; margin: 3px; display: inline-block;'>{cpc}</span>"
            for cpc in novel_cpcs])
        st.markdown(tags_html, unsafe_allow_html=True)

    st.divider()

    # Generated HyDE Abstract
    st.markdown("<h3 style='text-align: center;'>Generated HyDE Abstract</h3>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 0.9em; color: #666;'>HyDE stands for 'Hypothetical Document Embedding'. The AI writes this \"perfect\" abstract for a patent that matches your idea. It then converts this abstract into a vector in hopes to find the most semantically similar patents in the database.</p>", unsafe_allow_html=True)
    st.markdown("> This is the 'hypothetical patent' the AI used to search for semantic matches.")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Base Technology Search**")
        st.info(base_artifacts.get('hyde_abstract', 'N/A'))

    with col2:
        st.markdown("**Novel Features Search**")
        st.info(novel_artifacts.get('hyde_abstract', 'N/A'))


# --- Page Configuration ---
st.set_page_config(
    page_title="Patent Prior Art Analyst",
    page_icon="🤖",
    layout="wide"
)

# --- Rate Limiter Setup ---
# Initialize rate limiter: 5 queries per session
rate_limiter = SimpleRateLimiter(max_queries_per_session=5)

# --- Page Title & Description ---
st.title("Patent & Prior Art Analyst Agent")
st.markdown("""
Welcome! This application uses an advanced AI-powered agent to help you research prior art for your invention.
Simply describe your idea in plain English, and the agent will:

1.  **Transform** your idea into technical keywords, a hypothetical abstract, and relevant CPC codes.
2.  **Search** a patent database for semantically similar and relevant documents.
3.  **Synthesize** the findings into an expert-level summary.

📖 **Read more about this app and its limitations [HERE](https://github.com/rayvantsahni/Patent-and-Prior-Art-Analyst/blob/main/README.md).**
""")

st.divider()

# --- Usage Indicator ---
rate_limiter.show_usage_indicator()

# --- Usage Limits Info ---
with st.expander("ℹ️ About Usage Limits"):
    st.markdown("""
    **Why are there usage limits?**

    This app uses advanced AI models and vector databases which have associated costs.
    To keep the service free and available for everyone, I've implemented fair usage limits:

    - **5 queries per browser session** - Each session gets 5 free analyses
    - **Refresh for more** - You can refresh your browser to start a new session
    - **Please use responsibly** - This is a free service, so users are trusted to be fair

    **How it works:**
    - Your query counter is tied to your browser session
    - After 5 queries, you'll need to refresh the page to continue
    - This is a "soft limit" that relies on honest usage
    """)

st.divider()

# --- Input Area ---
user_description = st.text_area(
    "Describe your invention here:",
    height=200,
    placeholder="e.g., A smart coffee mug that uses a rechargeable battery to keep coffee at a perfect 140°F and "
                "syncs with a mobile app..."
)

# --- Button Styling (Safer Injection) ---
# Inject the custom CSS for the button
st.markdown(
    """
<style>
div.stButton > button:first-child {
    background-color: #4CAF50; /* Green */
    color: white;
    border: none;
    padding: 10px 20px;
    text-align: center;
    text-decoration: none;
    display: inline-block;
    font-size: 16px;
    margin: 4px 2px;
    cursor: pointer;
    border-radius: 8px;
}
div.stButton > button:first-child:hover {
    background-color: #45a049; /* Darker green */
}
</style>
""",
    unsafe_allow_html=True
)

# --- Analysis Button Logic ---
# Use a standard 'st.button' check, which is more reliable
if st.button("Analyze Prior Art"):
    if not user_description.strip():
        st.warning("Please enter a description of your invention.")
    elif not rate_limiter.can_query():
        st.error("⛔ Session limit reached (5 queries used). Please refresh your browser to start a new session with 5 fresh queries.")
    else:
        # Show a spinner while the agent is working
        with st.spinner("Agent is analyzing... This may take a moment."):
            try:
                # Call agent and get dictionary
                analysis_result = agent.run_analysis(user_description)

                # Check for errors
                if "error" in analysis_result:
                    st.error(analysis_result["error"])
                    st.stop()

                # Increment query counter after successful analysis
                rate_limiter.increment()

                final_report = analysis_result["final_report"]
                search_artifacts = analysis_result["search_artifacts"]

                # --- Display Final Report ---
                # st.header("Analysis Report")
                st.markdown("<h3 style='text-align: center; text-decoration: underline;'>Analysis Report</h3>",
                            unsafe_allow_html=True)
                st.markdown(final_report)

                # --- Display Intermediate Search Strategy ---
                with st.expander("See AI Search Strategy (Click to Open)"):
                    st.markdown(
                        "This is how the AI interpreted your idea to search the patent database. You can use these artifacts for your own research as well.")

                    base_artifacts = search_artifacts.get("base_technology_search", {})
                    novel_artifacts = search_artifacts.get("novel_features_search", {})

                    _display_artifacts_side_by_side(base_artifacts, novel_artifacts)

                # --- Copy Final Report ---
                with st.expander("Copy Report Text"):
                    st.code(final_report, language="markdown")

                # --- Download Buttons ---
                st.subheader("Download Report")

                # Generate a single, unique filename prefix (using updated function name)
                file_prefix = _generate_filename_prefix()

                # Use columns to place buttons side-by-side
                col1, col2, col3 = st.columns(3)

                with col1:
                    st.download_button(
                        label="Download as Text (.txt)",
                        data=final_report,
                        file_name=f"{file_prefix}.txt",
                        mime="text/plain"
                    )

                with col2:
                    st.download_button(
                        label="Download as Markdown (.md)",
                        data=final_report,
                        file_name=f"{file_prefix}.md",
                        mime="text/markdown"
                    )

                with col3:
                    # Create the docx file in memory (using updated function name)
                    docx_data = _create_docx(final_report)
                    st.download_button(
                        label="Download as Word (.docx)",
                        data=docx_data,
                        file_name=f"{file_prefix}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            except Exception as e:
                st.error(f"An error occurred during analysis: {e}")
